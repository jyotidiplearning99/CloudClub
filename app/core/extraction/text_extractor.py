"""
Extract text from PDF and DOCX files with HEADER/FOOTER support.
"""

import io
from pdfminer.high_level import extract_text as pdf_extract
from docx import Document
from app.utils.logger import get_logger

logger = get_logger(__name__)


class TextExtractor:
    """Extract text from various document formats."""
    
    def extract(self, file_bytes: bytes, filename: str) -> str:
        """
        Auto-detect format and extract text.
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename (for extension detection)
            
        Returns:
            Extracted text
        """
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            return self._extract_pdf(file_bytes)
        elif ext in ('docx', 'doc'):
            return self._extract_docx_with_headers(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX supported.")
    
    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file."""
        try:
            text = pdf_extract(io.BytesIO(file_bytes))
            logger.info("pdf_extracted", length=len(text))
            return text.strip()
        except Exception as e:
            logger.error("pdf_extraction_failed", error=str(e))
            raise ValueError(f"Failed to extract PDF: {e}")
    
    def _extract_docx_with_headers(self, file_bytes: bytes) -> str:
        """
        CRITICAL FIX: Extract DOCX with headers first (where name often is).
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            text_parts = []
            
            # CRITICAL: Extract headers FIRST (where name is)
            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                        logger.info("header_extracted", text=para.text.strip()[:50])
            
            # Extract body paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            # Extract footers
            for section in doc.sections:
                footer = section.footer
                for para in footer.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
            
            text = '\n'.join(text_parts)
            logger.info("docx_extracted_with_headers", length=len(text))
            return text.strip()
        except Exception as e:
            logger.error("docx_extraction_failed", error=str(e))
            raise ValueError(f"Failed to extract DOCX: {e}")
